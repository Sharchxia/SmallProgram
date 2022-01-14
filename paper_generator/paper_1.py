import tkinter as tk
from docx import Document
import random as rd
import tkinter.messagebox as msb

# 基本（全局）变量声明
w = tk.Tk()  # 主窗口
m = tk.Menu(w)  # 菜单
w.geometry("1100x700+200+50")
w.title('试卷')

fillProblem = []  # 缓存填空题题干和答案，如：[(题干，答案)...]
checkProblem = []  # 缓存判断题，如：[(题干，答案)...]
choiceProblem = []  # 缓存选择题，如：[(题干，[A，B，C，D]，答案)...]

Choice = []  # 缓存所涉选择题的答案
Fill = []  # 存储相应填空题输入框的变量
FillAnswer = []  # 缓存填空答案
Check = []  # 缓存判断答案

checkFlag = [None, '√', '×']
choiceFlag = [None, 'A', 'B', 'C', 'D']

choiceAnswer = []  # 存储选择题按键按钮变量
for v in range(4):
    choiceAnswer.append(tk.IntVar())
checkAnswer = []  # 存储判断题按键按钮变量
for v in range(2):
    checkAnswer.append(tk.IntVar())

# 三个组件框包含三种题型
fra_cho = tk.Frame(w, relief='ridge', borderwidth=5)
fra_che = tk.Frame(w, relief='ridge', borderwidth=5)
fra_fill = tk.Frame(w, relief='ridge', borderwidth=5)
fra_cho.place(relx=0, rely=0, relwidth=0.5, relheight=1)
fra_che.place(relx=0.5, rely=0, relwidth=0.5, relheight=0.35)
fra_fill.place(relx=0.5, rely=0.35, relwidth=0.5, relheight=.65)
w['menu'] = m  # 菜单绑定


def gen_pro_pool():  # 从word文档中提取出提
    # global fillProblem,checkProblem,choiceProblem
    doc = Document('test.docx')
    fillTable = doc.tables[0]  # 填空
    checkTable = doc.tables[1]  # 判断
    choiceTable = doc.tables[2]  # 选择
    element = []  # 用于将选择题4个选项组合的临时列表
    for i in range(8):
        fillProblem.append((fillTable.cell(i, 0).text, fillTable.cell(i, 1).text))
    for i in range(6):
        checkProblem.append((checkTable.cell(i, 0).text, checkTable.cell(i, 1).text))
    for i in range(8):
        for j in range(1, 5):
            element.append(choiceTable.cell(i, j).text)
        choiceProblem.append((choiceTable.cell(i, 0).text, element, choiceTable.cell(i, 5).text))
        element = []


def dec_ques_pool():  # 确定组卷用的题目（随机选择）
    num_title = []
    for j in range(3):  # 三种题
        num = []
        if j < 2:
            while len(num) < 4:  # 四个填空，四个选择
                i = rd.randint(0, 7)
                if i not in num:
                    num.append(i)
        else:
            while len(num) < 2:  # 两个判断
                i = rd.randint(0, 5)
                if i not in num:
                    num.append(i)
        num_title.append(num)  # 最终题目编号
    return num_title


def gen_gui():  # 界面可视化
    for i in range(len(finalProblem)):
        for j in finalProblem[i]:
            if i == 0:  # 选择可视化
                txt = tk.Text(fra_cho,font=18)
                txt.insert('1.0', str(finalProblem[i].index(j) + 1) + '. ' + choiceProblem[j][0])
                txt['state'] = tk.DISABLED
                txt.place(relx=0, relwidth=1, rely=finalProblem[i].index(j) * 0.25, relheight=0.05)
                for k in choiceProblem[j][1]:
                    rad = tk.Radiobutton(fra_cho, text=k, value=choiceProblem[j][1].index(k) + 1,selectcolor='#0ffff0',
                                         variable=choiceAnswer[finalProblem[i].index(j)],font=18)
                    rad.place(relx=0, relwidth=1,
                              rely=finalProblem[i].index(j) * .25 + (choiceProblem[j][1].index(k) + 1) * .05)
                Choice.append(choiceProblem[j][2])  # 存储对应选择的答案
            elif i == 1:  # 填空可视化
                txt = tk.Text(fra_fill,font=18,bg='#f0f000')
                txt.insert('1.0', str(finalProblem[i].index(j) + 7) + '. ' + fillProblem[j][0])
                txt['state'] = tk.DISABLED
                txt.place(relx=0, relwidth=1, rely=finalProblem[i].index(j) * 0.25, relheight=0.125)
                txtAnswer = tk.Text(fra_fill,font=22)
                Fill.append(txtAnswer)
                txtAnswer.place(relx=0, relwidth=1, rely=finalProblem[i].index(j) * 0.25 + 0.125, relheight=0.125)
                FillAnswer.append(fillProblem[j][1])  # 存储对应填空答案
            else:  # 判断可视化
                txt = tk.Text(fra_che,font=18)
                txt.insert('1.0', str(finalProblem[i].index(j) + 5) + '. ' + checkProblem[j][0])
                txt['state'] = tk.DISABLED
                txt.place(relx=0, relwidth=1, rely=finalProblem[i].index(j) * 0.5, relheight=0.2)
                rad = tk.Radiobutton(fra_che, text=checkFlag[1], value=1,selectcolor='#0ffff0',
                                     variable=checkAnswer[finalProblem[i].index(j)],font=18)
                rad.place(relx=0, relwidth=.5, rely=finalProblem[i].index(j) * 0.5 + 0.2, relheight=0.3)
                rad = tk.Radiobutton(fra_che, text=checkFlag[2], value=2,selectcolor='#0ffff0',
                                     variable=checkAnswer[finalProblem[i].index(j)],font=18)
                rad.place(relx=0.5, relwidth=.5, rely=finalProblem[i].index(j) * 0.5 + 0.2, relheight=0.3)
                Check.append(checkProblem[j][1])  # 存储对应判断答案


def correct():  # 改卷
    score = 0  # 分数
    wrong = []  # 错误题号
    for i in range(len(choiceAnswer)):
        if choiceFlag[choiceAnswer[i].get()] == Choice[i]:
            score += 10
        else:
            wrong.append(str(i + 1) + '题')
    for i in range(len(checkAnswer)):
        if checkFlag[checkAnswer[i].get()] == Check[i]:
            score += 10
        else:
            wrong.append(str(5 + i) + '题')
    for i in range(len(Fill)):
        if Fill[i].get(0.0, 'end').replace('\n','').replace(' ','') == FillAnswer[i]:
            score += 10
        else:
            wrong.append(str(i + 7) + '题')
    wrongInfo = '，'.join(wrong)
    msb.showinfo('得分及错题', '总分：' + str(score) + '\n错题：' + wrongInfo)  # 显示分数及错题


# 涉及函数调用如下
m.add_command(label='提交试卷', command=correct)
finalProblem = dec_ques_pool()
gen_pro_pool()
gen_gui()

tk.mainloop()

# print(fillProblem)
# print(choiceProblem)
# print(checkProblem)
# print(finalProblem)
